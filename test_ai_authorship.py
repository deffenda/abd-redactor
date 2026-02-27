import json
import importlib
from pathlib import Path
import sys
from types import SimpleNamespace

import fitz
import pytest

sys.path.insert(0, str(Path(__file__).parent / "src"))

from ai_authorship import analyze_file_authorship, analyze_text_authorship, extract_text_from_file
from upload_redaction import redact_uploaded_file


def _make_pdf(path: Path, text: str) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(path)
    doc.close()


def test_extract_text_from_pdf(tmp_path: Path) -> None:
    pdf_path = tmp_path / "sample.pdf"
    _make_pdf(pdf_path, "This is a sample sentence from a PDF document.")

    text = extract_text_from_file(pdf_path)

    assert "sample sentence" in text.lower()


def test_analyze_text_authorship_returns_metrics() -> None:
    text = (
        "The project team met on Monday to review the delivery timeline. "
        "We identified two blockers, assigned owners, and agreed on a mitigation plan."
    )

    result = analyze_text_authorship(text)

    assert 0.0 <= result.ai_probability <= 1.0
    assert 0.0 <= result.confidence <= 1.0
    assert result.features.word_count > 0
    assert result.label in {"likely_ai", "likely_human", "inconclusive"}


def test_extract_text_from_docx(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document

    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("This paragraph came from a Word document.")
    document.save(docx_path)

    text = extract_text_from_file(docx_path)

    assert "word document" in text.lower()


def test_analyze_endpoint_pdf(tmp_path: Path) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    from ai_authorship_bot import app

    pdf_path = tmp_path / "endpoint.pdf"
    _make_pdf(pdf_path, "The meeting notes were finalized and shared with the team.")

    client = TestClient(app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/analyze",
            data={"detector": "heuristic"},
            files={"file": ("endpoint.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"] == "endpoint.pdf"
    assert payload["detector_requested"] == "heuristic"
    assert "ai_probability_percent" in payload


def test_invalid_detector_for_file_analysis(tmp_path: Path) -> None:
    pdf_path = tmp_path / "invalid-detector.pdf"
    _make_pdf(pdf_path, "A short document for detector validation.")

    with pytest.raises(ValueError):
        analyze_file_authorship(pdf_path, detector="unsupported")


def test_analyze_endpoint_rejects_invalid_detector(tmp_path: Path) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    from ai_authorship_bot import app

    pdf_path = tmp_path / "invalid-endpoint.pdf"
    _make_pdf(pdf_path, "A short document for endpoint validation.")

    client = TestClient(app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/analyze",
            data={"detector": "not-a-detector"},
            files={"file": ("invalid-endpoint.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 400
    assert "Unsupported detector" in response.json()["detail"]


def test_model_detector_requires_api_key(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    from ai_authorship_bot import app

    monkeypatch.delenv("OPENAI_API_KEY", raising=False)

    pdf_path = tmp_path / "model-endpoint.pdf"
    _make_pdf(pdf_path, "A short document for model-detector validation.")

    client = TestClient(app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/analyze",
            data={"detector": "model"},
            files={"file": ("model-endpoint.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 503
    assert "OPENAI_API_KEY" in response.json()["detail"]


def test_analyze_endpoint_rejects_invalid_model_name(tmp_path: Path) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    from ai_authorship_bot import app

    pdf_path = tmp_path / "invalid-model.pdf"
    _make_pdf(pdf_path, "Text for invalid model test.")

    client = TestClient(app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/analyze",
            data={"detector": "model", "model_name": "not-a-real-model"},
            files={"file": ("invalid-model.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 400
    assert "Unsupported model" in response.json()["detail"]


def test_analyze_endpoint_passes_selected_model(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    import ai_authorship_bot

    captured: dict[str, str | None] = {}

    def _fake_analyze_file_authorship(path: Path, filename: str | None = None, detector: str = "heuristic", model_name: str | None = None):
        captured["detector"] = detector
        captured["model_name"] = model_name
        result = SimpleNamespace(
            label="inconclusive",
            ai_probability=0.5,
            confidence=0.5,
            summary="mocked",
            features=SimpleNamespace(
                word_count=10,
                sentence_count=2,
                avg_sentence_length=5.0,
                sentence_length_stddev=1.0,
                type_token_ratio=0.7,
                repeated_bigram_ratio=0.0,
                stopword_ratio=0.4,
                punctuation_density=0.05,
            ),
        )
        return result, "mock text"

    monkeypatch.setattr(ai_authorship_bot, "analyze_file_authorship", _fake_analyze_file_authorship)

    pdf_path = tmp_path / "selected-model.pdf"
    _make_pdf(pdf_path, "Text for selected model test.")

    client = TestClient(ai_authorship_bot.app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/analyze",
            data={"detector": "model", "model_name": "gpt-5"},
            files={"file": ("selected-model.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["model_name_requested"] == "gpt-5"
    assert captured["detector"] == "model"
    assert captured["model_name"] == "gpt-5"


def test_analyze_endpoint_accepts_bedrock_model_name(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    import ai_authorship_bot

    captured: dict[str, str | None] = {}

    def _fake_analyze_file_authorship(path: Path, filename: str | None = None, detector: str = "heuristic", model_name: str | None = None):
        _ = path
        _ = filename
        captured["detector"] = detector
        captured["model_name"] = model_name
        result = SimpleNamespace(
            label="inconclusive",
            ai_probability=0.5,
            confidence=0.5,
            summary="mocked",
            features=SimpleNamespace(
                word_count=10,
                sentence_count=2,
                avg_sentence_length=5.0,
                sentence_length_stddev=1.0,
                type_token_ratio=0.7,
                repeated_bigram_ratio=0.0,
                stopword_ratio=0.4,
                punctuation_density=0.05,
            ),
        )
        return result, "mock text"

    monkeypatch.setattr(ai_authorship_bot, "analyze_file_authorship", _fake_analyze_file_authorship)

    pdf_path = tmp_path / "selected-bedrock-model.pdf"
    _make_pdf(pdf_path, "Text for selected Bedrock model test.")

    selected_model = "bedrock:anthropic.claude-3-7-sonnet-20250219-v1:0"
    client = TestClient(ai_authorship_bot.app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/analyze",
            data={"detector": "model", "model_name": selected_model},
            files={"file": ("selected-bedrock-model.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["model_name_requested"] == selected_model
    assert captured["detector"] == "model"
    assert captured["model_name"] == selected_model


class _FakeComprehendClient:
    def detect_pii_entities(self, Text: str, LanguageCode: str) -> dict:
        entities = []
        for phrase, entity_type in (
            ("john.doe@example.com", "EMAIL"),
            ("123-45-6789", "SSN"),
            ("(555) 123-4567", "PHONE"),
        ):
            index = Text.find(phrase)
            if index >= 0:
                entities.append(
                    {
                        "Type": entity_type,
                        "Score": 0.99,
                        "BeginOffset": index,
                        "EndOffset": index + len(phrase),
                    }
                )
        return {"Entities": entities}


def test_redact_uploaded_pdf_uses_comprehend_pipeline(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("upload_redaction.comprehend_client", _FakeComprehendClient())
    pdf_path = tmp_path / "pii.pdf"
    _make_pdf(
        pdf_path,
        "Contact john.doe@example.com. SSN 123-45-6789. "
        "Call (555) 123-4567. Contact john.doe@example.com again.",
    )

    output_path, metrics = redact_uploaded_file(pdf_path, filename="pii.pdf")
    try:
        assert output_path.exists()
        assert metrics["total_findings_detected"] >= 3
        assert metrics["unique_phrases_detected"] >= 3
        assert metrics["total_redactions"] >= 1
        assert metrics["entities_by_type"].get("EMAIL", 0) >= 1
        with fitz.open(output_path) as doc:
            assert len(doc) >= 1
    finally:
        output_path.unlink(missing_ok=True)


def test_redact_endpoint_returns_download(tmp_path: Path) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    from ai_authorship_bot import app

    monkeypatch = pytest.MonkeyPatch()
    monkeypatch.setattr("upload_redaction.comprehend_client", _FakeComprehendClient())

    pdf_path = tmp_path / "upload.pdf"
    _make_pdf(pdf_path, "SSN 123-45-6789 Email john.doe@example.com")

    client = TestClient(app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/redact",
            files={"file": ("upload.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")
    assert "attachment;" in response.headers.get("content-disposition", "")
    assert "upload-redacted.pdf" in response.headers.get("content-disposition", "")
    assert "x-redaction-stats" in response.headers
    assert len(response.content) > 0
    monkeypatch.undo()


def test_redact_endpoint_rejects_invalid_engine(tmp_path: Path) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    from ai_authorship_bot import app

    pdf_path = tmp_path / "invalid-engine.pdf"
    _make_pdf(pdf_path, "Email john.doe@example.com")

    client = TestClient(app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/redact",
            data={"redaction_engine": "not-valid"},
            files={"file": ("invalid-engine.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 400
    assert "Unsupported redaction engine" in response.json()["detail"]


def test_redact_model_engine_requires_api_key(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    from ai_authorship_bot import app

    monkeypatch.delenv("OPENAI_API_KEY", raising=False)

    pdf_path = tmp_path / "model-redact.pdf"
    _make_pdf(pdf_path, "Email john.doe@example.com")

    client = TestClient(app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/redact",
            data={"redaction_engine": "model", "model_name": "gpt-5"},
            files={"file": ("model-redact.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 503
    assert "OPENAI_API_KEY" in response.json()["detail"]


def test_redact_endpoint_passes_selected_engine_and_model(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    import ai_authorship_bot

    captured: dict[str, str | None] = {}

    def _fake_redact_uploaded_file(input_path: Path, filename: str, **kwargs):
        captured["redaction_engine"] = kwargs.get("redaction_engine")
        captured["model_name"] = kwargs.get("model_name")
        output_path = tmp_path / "fake-redacted.pdf"
        _make_pdf(output_path, "Redacted")
        return output_path, {
            "model_name_used": kwargs.get("model_name"),
            "total_findings_detected": 0,
            "unique_phrases_detected": 0,
            "total_redactions": 0,
            "pages_with_redactions": 0,
            "redactions_per_page": {},
            "entities_by_type": {},
        }

    monkeypatch.setattr(ai_authorship_bot, "redact_uploaded_file", _fake_redact_uploaded_file)

    pdf_path = tmp_path / "selected-engine.pdf"
    _make_pdf(pdf_path, "Email john.doe@example.com")

    client = TestClient(ai_authorship_bot.app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/redact",
            data={"redaction_engine": "model", "model_name": "gpt-5"},
            files={"file": ("selected-engine.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 200
    assert captured["redaction_engine"] == "model"
    assert captured["model_name"] == "gpt-5"


def test_redact_endpoint_rejects_docx(tmp_path: Path) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    pytest.importorskip("docx")
    from docx import Document
    from fastapi.testclient import TestClient
    from ai_authorship_bot import app

    input_path = tmp_path / "upload.docx"
    document = Document()
    document.add_paragraph("PII content")
    document.save(input_path)

    client = TestClient(app)
    with input_path.open("rb") as handle:
        response = client.post(
            "/redact",
            files={
                "file": (
                    "upload.docx",
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 400
    assert "Supported redaction types: .pdf" in response.json()["detail"]


def test_document_summary_endpoint_returns_download_and_passes_model_and_directions(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    import ai_authorship_bot

    captured: dict[str, str | None] = {}

    def _fake_generate_document_summary_file(
        input_path: Path,
        *,
        filename: str,
        model_name: str | None = None,
        additional_directions: str | None = None,
    ):
        captured["filename"] = filename
        captured["model_name"] = model_name
        captured["additional_directions"] = additional_directions
        output_path = tmp_path / "document-summary.pdf"
        _make_pdf(output_path, "Document summary mocked output.")
        return output_path, {
            "model_name_used": model_name,
            "input_characters": 100,
            "characters_analyzed": 100,
            "key_point_count": 3,
            "additional_directions_provided": bool(additional_directions),
            "additional_directions_length": len((additional_directions or "").strip()),
        }

    monkeypatch.setattr(ai_authorship_bot, "generate_document_summary_file", _fake_generate_document_summary_file)

    pdf_path = tmp_path / "summary-input.pdf"
    _make_pdf(pdf_path, "Project timeline and risks were discussed with action owners assigned.")

    client = TestClient(ai_authorship_bot.app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/document-summary",
            data={"model_name": "gpt-5", "additional_directions": "Focus on risks and deadlines."},
            files={"file": ("summary-input.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")
    assert "attachment;" in response.headers.get("content-disposition", "")
    assert "summary-input-document-summary.pdf" in response.headers.get("content-disposition", "")
    assert "x-document-summary-stats" in response.headers
    stats = json.loads(response.headers["x-document-summary-stats"])
    assert stats["model_name_requested"] == "gpt-5"
    assert stats["key_point_count"] == 3
    assert stats["additional_directions_provided"] is True
    assert captured["model_name"] == "gpt-5"
    assert captured["additional_directions"] == "Focus on risks and deadlines."


def test_document_summary_endpoint_rejects_invalid_model_name(tmp_path: Path) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    from ai_authorship_bot import app

    pdf_path = tmp_path / "invalid-summary-model.pdf"
    _make_pdf(pdf_path, "Summary model validation text.")

    client = TestClient(app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/document-summary",
            data={"model_name": "not-a-real-model"},
            files={"file": ("invalid-summary-model.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 400
    assert "Unsupported model" in response.json()["detail"]


def test_document_summary_endpoint_rejects_invalid_file_type(tmp_path: Path) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    from ai_authorship_bot import app

    txt_path = tmp_path / "not-supported.txt"
    txt_path.write_text("plain text input", encoding="utf-8")

    client = TestClient(app)
    with txt_path.open("rb") as handle:
        response = client.post(
            "/document-summary",
            data={"model_name": "gpt-5"},
            files={"file": ("not-supported.txt", handle, "text/plain")},
        )

    assert response.status_code == 400
    assert "Supported document summary types" in response.json()["detail"]


def test_document_summary_requires_api_key(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    from ai_authorship_bot import app

    monkeypatch.delenv("OPENAI_API_KEY", raising=False)

    pdf_path = tmp_path / "document-summary-api-key.pdf"
    _make_pdf(pdf_path, "Summary text for API key validation.")

    client = TestClient(app)
    with pdf_path.open("rb") as handle:
        response = client.post(
            "/document-summary",
            data={"model_name": "gpt-5"},
            files={"file": ("document-summary-api-key.pdf", handle, "application/pdf")},
        )

    assert response.status_code == 503
    assert "OPENAI_API_KEY" in response.json()["detail"]


def test_document_summary_uses_chunking_pipeline(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    import document_summary

    pdf_path = tmp_path / "chunked-summary.pdf"
    _make_pdf(
        pdf_path,
        "Alpha details. " * 80 + "Beta details. " * 80 + "Gamma details. " * 80,
    )

    observed_chunks: list[str] = []

    def _fake_chunk_text(_text: str, _max_chars: int) -> list[str]:
        return ["chunk-one-text", "chunk-two-text"]

    def _fake_call_document_summary_model(text: str, _model_name: str, additional_directions: str | None = None) -> dict:
        observed_chunks.append(text)
        return {
            "summary": f"Summary for {text}.",
            "key_points": [f"Point from {text}"],
            "action_items": [f"Action from {text}"],
            "relevant_details": [f"Detail from {text}"],
            "unanswered_questions": [],
            "disclaimer": f"Directions={additional_directions or ''}",
        }

    monkeypatch.setattr(document_summary, "chunk_text", _fake_chunk_text)
    monkeypatch.setattr(document_summary, "_call_document_summary_model", _fake_call_document_summary_model)

    output_path, metrics = document_summary.generate_document_summary_file(
        pdf_path,
        filename="chunked-summary.pdf",
        model_name="gpt-5",
        additional_directions="Focus on actions.",
    )
    try:
        assert output_path.exists()
        assert metrics["total_chunks"] == 2
        assert metrics["chunks_analyzed"] == 2
        assert metrics["characters_analyzed"] > 0
        assert observed_chunks == ["chunk-one-text", "chunk-two-text"]
    finally:
        output_path.unlink(missing_ok=True)


def test_s3_trigger_path_runs_all_three_capabilities(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("AWS_DEFAULT_REGION", "us-east-1")
    monkeypatch.setenv("AWS_ACCESS_KEY_ID", "test")
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "test")
    import assemble_output

    input_pdf = tmp_path / "incoming-s3.pdf"
    _make_pdf(input_pdf, "S3-trigger text for redaction, authorship, and summary.")

    uploaded_files: dict[tuple[str, str], dict[str, str]] = {}
    uploaded_json: dict[tuple[str, str], dict] = {}

    class _FakeS3Client:
        def upload_file(self, filename: str, bucket: str, key: str, ExtraArgs: dict | None = None) -> None:
            uploaded_files[(bucket, key)] = {
                "filename": filename,
                "content_type": (ExtraArgs or {}).get("ContentType", ""),
            }

    def _fake_upload_json(bucket: str, key: str, payload: dict) -> None:
        uploaded_json[(bucket, key)] = payload

    def _fake_analyze_file_authorship(
        _path: Path,
        filename: str | None = None,
        detector: str = "heuristic",
        model_name: str | None = None,
    ):
        _ = filename
        _ = detector
        _ = model_name
        result = SimpleNamespace(
            label="likely_human",
            ai_probability=0.2,
            confidence=0.9,
            summary="Mocked authorship summary.",
            features=SimpleNamespace(
                word_count=120,
                sentence_count=8,
                avg_sentence_length=15.0,
                sentence_length_stddev=4.0,
                type_token_ratio=0.7,
                repeated_bigram_ratio=0.01,
                stopword_ratio=0.4,
                punctuation_density=0.05,
            ),
        )
        return result, "mocked extracted text"

    def _fake_generate_document_summary_file(
        _input_path: Path,
        *,
        filename: str,
        model_name: str | None = None,
        additional_directions: str | None = None,
    ):
        _ = filename
        _ = model_name
        _ = additional_directions
        output_path = tmp_path / "s3-summary.pdf"
        _make_pdf(output_path, "Mocked S3 document summary.")
        return output_path, {
            "model_name_used": model_name or "gpt-5",
            "input_characters": 300,
            "characters_analyzed": 300,
            "total_chunks": 3,
            "chunks_analyzed": 3,
            "chunk_char_limit": 4500,
            "max_chunks": 12,
            "key_point_count": 4,
            "additional_directions_provided": bool(additional_directions),
            "additional_directions_length": len((additional_directions or "").strip()),
        }

    monkeypatch.setattr(assemble_output, "s3_client", _FakeS3Client())
    monkeypatch.setattr(assemble_output, "_upload_json", _fake_upload_json)
    monkeypatch.setattr(assemble_output, "analyze_file_authorship", _fake_analyze_file_authorship)
    monkeypatch.setattr(assemble_output, "generate_document_summary_file", _fake_generate_document_summary_file)

    outputs = assemble_output._run_s3_capabilities(
        local_input=input_pdf,
        input_key="incoming/sample.pdf",
        input_prefix="incoming/",
        output_bucket="example-bucket",
        report_prefix="reports/",
        include_authorship=True,
        include_document_summary=True,
        require_all_capabilities=True,
        authorship_detector="heuristic",
        authorship_model_name=None,
        document_summary_model_name="gpt-5",
        document_summary_directions="Focus on risks.",
    )

    assert outputs["authorship"]["status"] == "COMPLETED"
    assert outputs["document_summary"]["status"] == "COMPLETED"
    assert ("example-bucket", "reports/sample-authorship-report.json") in uploaded_json
    assert ("example-bucket", "reports/sample-document-summary.pdf") in uploaded_files


def test_modular_s3_pipeline_finalizes_all_capabilities(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("AWS_DEFAULT_REGION", "us-east-1")
    monkeypatch.setenv("AWS_ACCESS_KEY_ID", "test")
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "test")
    import assemble_output

    input_pdf = tmp_path / "incoming-modular.pdf"
    _make_pdf(input_pdf, "Modular pipeline source text.")

    uploaded_files: dict[tuple[str, str], dict[str, str]] = {}
    uploaded_json: dict[tuple[str, str], dict] = {}

    class _FakeS3Client:
        def download_file(self, _bucket: str, _key: str, filename: str) -> None:
            Path(filename).write_bytes(input_pdf.read_bytes())

        def upload_file(self, filename: str, bucket: str, key: str, ExtraArgs: dict | None = None) -> None:
            uploaded_files[(bucket, key)] = {
                "filename": filename,
                "content_type": (ExtraArgs or {}).get("ContentType", ""),
            }

    def _fake_upload_json(bucket: str, key: str, payload: dict) -> None:
        uploaded_json[(bucket, key)] = payload

    def _fake_load_json(bucket: str, key: str) -> dict:
        return uploaded_json[(bucket, key)]

    def _fake_analyze_file_authorship(
        _path: Path,
        filename: str | None = None,
        detector: str = "heuristic",
        model_name: str | None = None,
    ):
        _ = filename
        _ = detector
        _ = model_name
        result = SimpleNamespace(
            label="likely_human",
            ai_probability=0.22,
            confidence=0.88,
            summary="Authorship summary.",
            features=SimpleNamespace(
                word_count=100,
                sentence_count=7,
                avg_sentence_length=14.2,
                sentence_length_stddev=3.1,
                type_token_ratio=0.68,
                repeated_bigram_ratio=0.02,
                stopword_ratio=0.42,
                punctuation_density=0.06,
            ),
        )
        return result, "modular extracted text"

    def _fake_generate_document_summary_file(
        _input_path: Path,
        *,
        filename: str,
        model_name: str | None = None,
        additional_directions: str | None = None,
    ):
        _ = filename
        _ = model_name
        _ = additional_directions
        output_path = tmp_path / "modular-summary.pdf"
        _make_pdf(output_path, "Modular summary artifact.")
        return output_path, {
            "model_name_used": model_name or "gpt-5",
            "input_characters": 120,
            "characters_analyzed": 120,
            "total_chunks": 1,
            "chunks_analyzed": 1,
            "chunk_char_limit": 4500,
            "max_chunks": 12,
            "key_point_count": 2,
            "additional_directions_provided": bool(additional_directions),
            "additional_directions_length": len((additional_directions or "").strip()),
        }

    monkeypatch.setattr(assemble_output, "s3_client", _FakeS3Client())
    monkeypatch.setattr(assemble_output, "_upload_json", _fake_upload_json)
    monkeypatch.setattr(assemble_output, "_load_json", _fake_load_json)
    monkeypatch.setattr(assemble_output, "_delete_prefix", lambda _bucket, _prefix: 4)
    monkeypatch.setattr(assemble_output, "analyze_file_authorship", _fake_analyze_file_authorship)
    monkeypatch.setattr(assemble_output, "generate_document_summary_file", _fake_generate_document_summary_file)

    base_event = {
        "job_id": "job-modular-123",
        "input_bucket": "input-bucket",
        "input_key": "incoming/sample.pdf",
        "output_bucket": "output-bucket",
        "work_bucket": "output-bucket",
        "input_prefix": "incoming/",
        "report_prefix": "reports/",
        "work_prefix": "work/",
        "redacted_pdf_key": "redacted/sample-redacted.pdf",
        "report_key": "reports/sample-redaction-report.json",
        "base_report_key": "work/job-modular-123/base-report.json",
        "changes_made": 18,
        "effectiveness_score": 97.5,
        "enable_s3_authorship": True,
        "enable_s3_document_summary": True,
        "require_s3_capabilities": True,
        "s3_authorship_detector": "heuristic",
        "s3_authorship_model_name": "",
        "s3_summary_model_name": "gpt-5",
        "s3_summary_directions": "Highlight decisions.",
    }
    uploaded_json[("output-bucket", "work/job-modular-123/base-report.json")] = {
        "job_id": "job-modular-123",
        "output": {
            "bucket": "output-bucket",
            "redacted_pdf_key": "redacted/sample-redacted.pdf",
            "report_key": "reports/sample-redaction-report.json",
        },
    }

    authorship_output = assemble_output.authorship_artifact_lambda_handler(base_event, None)
    summary_output = assemble_output.document_summary_artifact_lambda_handler(base_event, None)
    final_result = assemble_output.finalize_report_lambda_handler(
        {
            **base_event,
            "capability_results": [authorship_output, summary_output],
        },
        None,
    )

    assert final_result["status"] == "COMPLETED"
    assert final_result["capabilities"]["authorship"]["status"] == "COMPLETED"
    assert final_result["capabilities"]["document_summary"]["status"] == "COMPLETED"
    assert final_result["authorship_report_key"] == "reports/sample-authorship-report.json"
    assert final_result["document_summary_pdf_key"] == "reports/sample-document-summary.pdf"
    assert ("output-bucket", "reports/sample-authorship-report.json") in uploaded_json
    assert ("output-bucket", "reports/sample-document-summary.pdf") in uploaded_files
    assert ("output-bucket", "reports/sample-redaction-report.json") in uploaded_json


def test_start_execution_includes_s3_capability_settings(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("AWS_DEFAULT_REGION", "us-east-1")
    monkeypatch.setenv("AWS_ACCESS_KEY_ID", "test")
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "test")
    monkeypatch.setenv("ENABLE_S3_AUTHORSHIP", "true")
    monkeypatch.setenv("ENABLE_S3_DOCUMENT_SUMMARY", "true")
    monkeypatch.setenv("REQUIRE_S3_CAPABILITIES", "false")
    monkeypatch.setenv("S3_AUTHORSHIP_DETECTOR", "model")
    monkeypatch.setenv("S3_AUTHORSHIP_MODEL", "gpt-5")
    monkeypatch.setenv("S3_SUMMARY_MODEL", "gpt-5")
    monkeypatch.setenv("S3_SUMMARY_DIRECTIONS", "Highlight deadlines.")

    import start_execution

    reloaded = importlib.reload(start_execution)
    payload = reloaded._build_execution_input("bucket-a", "incoming/sample.pdf")

    assert payload["enable_s3_authorship"] is True
    assert payload["enable_s3_document_summary"] is True
    assert payload["require_s3_capabilities"] is False
    assert payload["s3_authorship_detector"] == "model"
    assert payload["s3_authorship_model_name"] == "gpt-5"
    assert payload["s3_summary_model_name"] == "gpt-5"
    assert payload["s3_summary_directions"] == "Highlight deadlines."
